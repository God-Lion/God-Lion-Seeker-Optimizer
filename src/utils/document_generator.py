"""
Document Generator Module
Generates professional CV/Resume documents in both PDF and DOCX formats.
"""

from typing import List, Dict, Optional, Literal
from pathlib import Path
from reportlab.lib.pagesizes import letter
from reportlab.lib.units import inch
from reportlab.pdfgen import canvas
from reportlab.lib.colors import HexColor
from reportlab.pdfbase import pdfmetrics
from reportlab.pdfbase.ttfonts import TTFont
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.shared import OxmlElement
from docx.oxml.ns import qn


class DocumentGenerator:
    """
    A unified document generator for creating professional CVs/Resumes
    in both PDF and DOCX formats with consistent styling.
    """
    
    # Color definitions
    COLORS = {
        "text": "#000000",
        "link": "#0563C1",
        "heading": "#000000"
    }
    
    # Font sizes
    FONT_SIZES = {
        "name": 24,
        "title": 18,
        "heading": 14,
        "job_title": 12,
        "body": 11,
        "small": 10
    }
    
    def __init__(self, output_format: Literal["pdf", "docx", "both"] = "both"):
        """
        Initialize the document generator.
        
        Args:
            output_format: Output format - "pdf", "docx", or "both"
        """
        self.output_format = output_format
        self.pdf_canvas = None
        self.docx_doc = None
        self.current_y = None
        
    def _hex_to_rgb(self, hex_color: str) -> tuple:
        """Convert hex color to RGB tuple."""
        hex_color = hex_color.lstrip('#')
        return tuple(int(hex_color[i:i+2], 16) for i in (0, 2, 4))
    
    def _init_pdf(self, filename: str):
        """Initialize PDF canvas."""
        self.pdf_canvas = canvas.Canvas(filename, pagesize=letter)
        self.width, self.height = letter
        self.margins = {
            "top": self.height - 36,
            "bottom": 36,
            "left": 72,
            "right": self.width - 72
        }
        self.current_y = self.margins["top"]
        
        # Try to register custom fonts (fallback to standard fonts if unavailable)
        try:
            pdfmetrics.registerFont(TTFont('Calibri', 'calibri.ttf'))
            pdfmetrics.registerFont(TTFont('Calibri-Bold', 'calibrib.ttf'))
        except:
            pass
    
    def _init_docx(self):
        """Initialize DOCX document."""
        self.docx_doc = Document()
        
        # Set document margins (in inches)
        sections = self.docx_doc.sections
        for section in sections:
            section.top_margin = Inches(0.5)
            section.bottom_margin = Inches(0.5)
            section.left_margin = Inches(1.0)
            section.right_margin = Inches(1.0)
    
    def _add_hyperlink_docx(self, paragraph, text: str, url: str):
        """Add a hyperlink to a DOCX paragraph."""
        part = paragraph.part
        r_id = part.relate_to(url, "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink", is_external=True)
        
        hyperlink = OxmlElement('w:hyperlink')
        hyperlink.set(qn('r:id'), r_id)
        
        new_run = OxmlElement('w:r')
        rPr = OxmlElement('w:rPr')
        
        # Style as link (blue, underline)
        c = OxmlElement('w:color')
        c.set(qn('w:val'), '0563C1')
        rPr.append(c)
        
        u = OxmlElement('w:u')
        u.set(qn('w:val'), 'single')
        rPr.append(u)
        
        new_run.append(rPr)
        new_run.text = text
        hyperlink.append(new_run)
        
        paragraph._element.append(hyperlink)
        
        return hyperlink
    
    def add_name(self, name: str):
        """Add name at top center."""
        if self.pdf_canvas:
            try:
                font = "Calibri-Bold"
                self.pdf_canvas.setFont(font, self.FONT_SIZES["name"])
            except:
                font = "Helvetica-Bold"
                self.pdf_canvas.setFont(font, self.FONT_SIZES["name"])
            
            self.pdf_canvas.setFillColor(HexColor(self.COLORS["text"]))
            text_width = self.pdf_canvas.stringWidth(name, font, self.FONT_SIZES["name"])
            x = (self.width - text_width) / 2
            self.pdf_canvas.drawString(x, 756, name)
        
        if self.docx_doc:
            paragraph = self.docx_doc.add_paragraph()
            run = paragraph.add_run(name)
            run.font.size = Pt(self.FONT_SIZES["name"])
            run.font.bold = True
            run.font.name = 'Calibri'
            rgb = self._hex_to_rgb(self.COLORS["text"])
            run.font.color.rgb = RGBColor(*rgb)
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            paragraph.paragraph_format.space_after = Pt(6)
    
    def add_title(self, title: str):
        """Add professional title."""
        if self.pdf_canvas:
            try:
                font = "Calibri"
                self.pdf_canvas.setFont(font, self.FONT_SIZES["title"])
            except:
                font = "Helvetica"
                self.pdf_canvas.setFont(font, self.FONT_SIZES["title"])
            
            text_width = self.pdf_canvas.stringWidth(title, font, self.FONT_SIZES["title"])
            x = (self.width - text_width) / 2
            self.pdf_canvas.drawString(x, 743, title)
        
        if self.docx_doc:
            paragraph = self.docx_doc.add_paragraph()
            run = paragraph.add_run(title)
            run.font.size = Pt(self.FONT_SIZES["title"])
            run.font.name = 'Calibri'
            rgb = self._hex_to_rgb(self.COLORS["text"])
            run.font.color.rgb = RGBColor(*rgb)
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            paragraph.paragraph_format.space_after = Pt(12)
    
    def add_contact_info(self, email: str, linkedin: Optional[str] = None, 
                        github: Optional[str] = None, phone: Optional[str] = None):
        """Add contact information with links."""
        if self.pdf_canvas:
            y = 717
            try:
                font = "Calibri"
                self.pdf_canvas.setFont(font, self.FONT_SIZES["small"])
            except:
                font = "Helvetica"
                self.pdf_canvas.setFont(font, self.FONT_SIZES["small"])
            
            # Build contact string
            contact_parts = []
            contact_links = []
            
            if email:
                contact_parts.append(email)
                contact_links.append(("email", email, f"mailto:{email}"))
            
            if phone:
                contact_parts.append(phone)
                contact_links.append(("phone", phone, None))
            
            if linkedin:
                contact_parts.append("LinkedIn")
                contact_links.append(("linkedin", "LinkedIn", linkedin))
            
            if github:
                contact_parts.append("GitHub")
                contact_links.append(("github", "GitHub", github))
            
            # Calculate total width and starting position
            full_text = " | ".join(contact_parts)
            total_width = self.pdf_canvas.stringWidth(full_text, font, self.FONT_SIZES["small"])
            x_start = (self.width - total_width) / 2
            
            # Draw each part
            current_x = x_start
            for i, (link_type, text, url) in enumerate(contact_links):
                if i > 0:
                    self.pdf_canvas.setFillColor(HexColor(self.COLORS["text"]))
                    sep_text = " | "
                    self.pdf_canvas.drawString(current_x, y, sep_text)
                    current_x += self.pdf_canvas.stringWidth(sep_text, font, self.FONT_SIZES["small"])
                
                if url and url.startswith(("http", "mailto")):
                    self.pdf_canvas.setFillColor(HexColor(self.COLORS["link"]))
                else:
                    self.pdf_canvas.setFillColor(HexColor(self.COLORS["text"]))
                
                self.pdf_canvas.drawString(current_x, y, text)
                
                if url:
                    text_width = self.pdf_canvas.stringWidth(text, font, self.FONT_SIZES["small"])
                    self.pdf_canvas.linkURL(url, (current_x, y-3, current_x+text_width, y+13), relative=0)
                    current_x += text_width
                else:
                    current_x += self.pdf_canvas.stringWidth(text, font, self.FONT_SIZES["small"])
        
        if self.docx_doc:
            paragraph = self.docx_doc.add_paragraph()
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            contact_items = []
            if email:
                contact_items.append(("email", email, f"mailto:{email}"))
            if phone:
                contact_items.append(("phone", phone, None))
            if linkedin:
                contact_items.append(("linkedin", "LinkedIn", linkedin))
            if github:
                contact_items.append(("github", "GitHub", github))
            
            for i, (item_type, text, url) in enumerate(contact_items):
                if i > 0:
                    run = paragraph.add_run(" | ")
                    run.font.size = Pt(self.FONT_SIZES["small"])
                    run.font.name = 'Calibri'
                
                if url:
                    self._add_hyperlink_docx(paragraph, text, url)
                else:
                    run = paragraph.add_run(text)
                    run.font.size = Pt(self.FONT_SIZES["small"])
                    run.font.name = 'Calibri'
            
            paragraph.paragraph_format.space_after = Pt(18)
    
    def add_section_heading(self, heading: str, y_position: Optional[float] = None):
        """Add section heading with underline."""
        if self.pdf_canvas:
            if y_position is None:
                y_position = self.current_y - 30
            
            try:
                font = "Calibri-Bold"
                self.pdf_canvas.setFont(font, self.FONT_SIZES["heading"])
            except:
                font = "Helvetica-Bold"
                self.pdf_canvas.setFont(font, self.FONT_SIZES["heading"])
            
            self.pdf_canvas.setFillColor(HexColor(self.COLORS["heading"]))
            self.pdf_canvas.drawString(self.margins["left"], y_position, heading)
            
            # Underline
            text_width = self.pdf_canvas.stringWidth(heading, font, self.FONT_SIZES["heading"])
            self.pdf_canvas.line(self.margins["left"], y_position - 2, self.margins["left"] + text_width, y_position - 2)
            
            self.current_y = y_position - 20
            return self.current_y
        
        if self.docx_doc:
            paragraph = self.docx_doc.add_paragraph()
            run = paragraph.add_run(heading)
            run.font.size = Pt(self.FONT_SIZES["heading"])
            run.font.bold = True
            run.font.name = 'Calibri'
            rgb = self._hex_to_rgb(self.COLORS["heading"])
            run.font.color.rgb = RGBColor(*rgb)
            run.underline = True
            paragraph.paragraph_format.space_before = Pt(12)
            paragraph.paragraph_format.space_after = Pt(8)
    
    def add_paragraph(self, text: str, indent: bool = False):
        """Add a paragraph of text."""
        if self.pdf_canvas:
            try:
                font = "Calibri"
                self.pdf_canvas.setFont(font, self.FONT_SIZES["body"])
            except:
                font = "Helvetica"
                self.pdf_canvas.setFont(font, self.FONT_SIZES["body"])
            
            x = self.margins["left"] + (18 if indent else 0)
            max_width = self.margins["right"] - x
            
            # Word wrap
            words = text.split()
            lines = []
            current_line = []
            
            for word in words:
                test_line = ' '.join(current_line + [word])
                if self.pdf_canvas.stringWidth(test_line, font, self.FONT_SIZES["body"]) <= max_width:
                    current_line.append(word)
                else:
                    if current_line:
                        lines.append(' '.join(current_line))
                    current_line = [word]
            
            if current_line:
                lines.append(' '.join(current_line))
            
            for line in lines:
                self.pdf_canvas.drawString(x, self.current_y, line)
                self.current_y -= 15
            
            self.current_y -= 5
        
        if self.docx_doc:
            paragraph = self.docx_doc.add_paragraph()
            run = paragraph.add_run(text)
            run.font.size = Pt(self.FONT_SIZES["body"])
            run.font.name = 'Calibri'
            
            if indent:
                paragraph.paragraph_format.left_indent = Inches(0.25)
            
            paragraph.paragraph_format.space_after = Pt(6)
    
    def add_job_entry(self, title: str, company: str, location: Optional[str] = None,
                     dates: str = "", bullets: Optional[List[str]] = None, 
                     y_start: Optional[float] = None):
        """Add job experience entry."""
        if self.pdf_canvas:
            if y_start is None:
                y_start = self.current_y
            
            try:
                font_bold = "Calibri-Bold"
                font_regular = "Calibri"
                self.pdf_canvas.setFont(font_bold, self.FONT_SIZES["job_title"])
            except:
                font_bold = "Helvetica-Bold"
                font_regular = "Helvetica"
                self.pdf_canvas.setFont(font_bold, self.FONT_SIZES["job_title"])
            
            self.pdf_canvas.drawString(self.margins["left"], y_start, title)
            
            # Dates (right-aligned)
            self.pdf_canvas.setFont(font_regular, self.FONT_SIZES["body"])
            date_width = self.pdf_canvas.stringWidth(dates, font_regular, self.FONT_SIZES["body"])
            self.pdf_canvas.drawString(self.margins["right"] - date_width, y_start, dates)
            
            # Company and location
            y_start -= 15
            company_text = company
            if location:
                company_text += f" | {location}"
            self.pdf_canvas.setFont(font_regular, self.FONT_SIZES["body"])
            self.pdf_canvas.drawString(self.margins["left"], y_start, company_text)
            
            # Bullet points
            if bullets:
                y_start -= 20
                for bullet in bullets:
                    self.pdf_canvas.drawString(self.margins["left"], y_start, "•")
                    
                    # Word wrap bullet text
                    max_width = self.margins["right"] - (self.margins["left"] + 18)
                    words = bullet.split()
                    lines = []
                    current_line = []
                    
                    for word in words:
                        test_line = ' '.join(current_line + [word])
                        if self.pdf_canvas.stringWidth(test_line, font_regular, self.FONT_SIZES["body"]) <= max_width:
                            current_line.append(word)
                        else:
                            if current_line:
                                lines.append(' '.join(current_line))
                            current_line = [word]
                    
                    if current_line:
                        lines.append(' '.join(current_line))
                    
                    for line in lines:
                        x = self.margins["left"] + 18
                        self.pdf_canvas.drawString(x, y_start, line)
                        y_start -= 15
            
            self.current_y = y_start - 5
            return self.current_y
        
        if self.docx_doc:
            paragraph = self.docx_doc.add_paragraph()
            run = paragraph.add_run(title)
            run.font.size = Pt(self.FONT_SIZES["job_title"])
            run.font.bold = True
            run.font.name = 'Calibri'
            
            run = paragraph.add_run(f"\t{dates}")
            run.font.size = Pt(self.FONT_SIZES["body"])
            run.font.name = 'Calibri'
            paragraph.paragraph_format.space_after = Pt(2)
            
            paragraph = self.docx_doc.add_paragraph()
            company_text = company
            if location:
                company_text += f" | {location}"
            run = paragraph.add_run(company_text)
            run.font.size = Pt(self.FONT_SIZES["body"])
            run.font.italic = True
            run.font.name = 'Calibri'
            paragraph.paragraph_format.space_after = Pt(6)
            
            if bullets:
                for bullet in bullets:
                    paragraph = self.docx_doc.add_paragraph(bullet, style='List Bullet')
                    paragraph.paragraph_format.left_indent = Inches(0.25)
                    run = paragraph.runs[0]
                    run.font.size = Pt(self.FONT_SIZES["body"])
                    run.font.name = 'Calibri'
                
                paragraph.paragraph_format.space_after = Pt(12)
    
    def add_education_entry(self, degree: str, institution: str, 
                           location: Optional[str] = None, dates: str = "",
                           details: Optional[List[str]] = None,
                           y_start: Optional[float] = None):
        """Add education entry."""
        if self.pdf_canvas:
            if y_start is None:
                y_start = self.current_y
            
            try:
                font_bold = "Calibri-Bold"
                font_regular = "Calibri"
                self.pdf_canvas.setFont(font_bold, self.FONT_SIZES["job_title"])
            except:
                font_bold = "Helvetica-Bold"
                font_regular = "Helvetica"
                self.pdf_canvas.setFont(font_bold, self.FONT_SIZES["job_title"])
            
            self.pdf_canvas.drawString(self.margins["left"], y_start, degree)
            
            self.pdf_canvas.setFont(font_regular, self.FONT_SIZES["body"])
            date_width = self.pdf_canvas.stringWidth(dates, font_regular, self.FONT_SIZES["body"])
            self.pdf_canvas.drawString(self.margins["right"] - date_width, y_start, dates)
            
            y_start -= 15
            institution_text = institution
            if location:
                institution_text += f" | {location}"
            self.pdf_canvas.drawString(self.margins["left"], y_start, institution_text)
            
            if details:
                y_start -= 20
                for detail in details:
                    self.pdf_canvas.drawString(self.margins["left"], y_start, "•")
                    self.pdf_canvas.drawString(self.margins["left"] + 18, y_start, detail)
                    y_start -= 15
            
            self.current_y = y_start - 5
            return self.current_y
        
        if self.docx_doc:
            paragraph = self.docx_doc.add_paragraph()
            run = paragraph.add_run(degree)
            run.font.size = Pt(self.FONT_SIZES["job_title"])
            run.font.bold = True
            run.font.name = 'Calibri'
            
            run = paragraph.add_run(f"\t{dates}")
            run.font.size = Pt(self.FONT_SIZES["body"])
            run.font.name = 'Calibri'
            paragraph.paragraph_format.space_after = Pt(2)
            
            paragraph = self.docx_doc.add_paragraph()
            institution_text = institution
            if location:
                institution_text += f" | {location}"
            run = paragraph.add_run(institution_text)
            run.font.size = Pt(self.FONT_SIZES["body"])
            run.font.italic = True
            run.font.name = 'Calibri'
            
            if details:
                paragraph.paragraph_format.space_after = Pt(6)
                for detail in details:
                    para = self.docx_doc.add_paragraph(detail, style='List Bullet')
                    para.paragraph_format.left_indent = Inches(0.25)
                    run = para.runs[0]
                    run.font.size = Pt(self.FONT_SIZES["body"])
                    run.font.name = 'Calibri'
            else:
                paragraph.paragraph_format.space_after = Pt(12)
    
    def add_certification(self, name: str, issuer: str, date: Optional[str] = None,
                         credential_url: Optional[str] = None):
        """Add certification entry."""
        cert_text = f"{name} - {issuer}"
        if date:
            cert_text += f" ({date})"
        
        if self.pdf_canvas:
            try:
                font = "Calibri"
                self.pdf_canvas.setFont(font, self.FONT_SIZES["body"])
            except:
                font = "Helvetica"
                self.pdf_canvas.setFont(font, self.FONT_SIZES["body"])
            
            self.pdf_canvas.drawString(self.margins["left"], self.current_y, "•")
            
            if credential_url:
                self.pdf_canvas.setFillColor(HexColor(self.COLORS["link"]))
                self.pdf_canvas.drawString(self.margins["left"] + 18, self.current_y, cert_text)
                text_width = self.pdf_canvas.stringWidth(cert_text, font, self.FONT_SIZES["body"])
                self.pdf_canvas.linkURL(credential_url,
                                      (self.margins["left"] + 18, self.current_y - 3,
                                       self.margins["left"] + 18 + text_width, self.current_y + 13),
                                      relative=0)
                self.pdf_canvas.setFillColor(HexColor(self.COLORS["text"]))
            else:
                self.pdf_canvas.drawString(self.margins["left"] + 18, self.current_y, cert_text)
            
            self.current_y -= 15
        
        if self.docx_doc:
            paragraph = self.docx_doc.add_paragraph(style='List Bullet')
            paragraph.paragraph_format.left_indent = Inches(0.25)
            
            if credential_url:
                self._add_hyperlink_docx(paragraph, cert_text, credential_url)
            else:
                run = paragraph.add_run(cert_text)
                run.font.size = Pt(self.FONT_SIZES["body"])
                run.font.name = 'Calibri'
    
    def add_skills_section(self, skills_dict: Dict[str, List[str]]):
        """Add skills section with categories."""
        if self.pdf_canvas:
            try:
                font_bold = "Calibri-Bold"
                font_regular = "Calibri"
            except:
                font_bold = "Helvetica-Bold"
                font_regular = "Helvetica"
            
            for category, skills in skills_dict.items():
                self.pdf_canvas.setFont(font_bold, self.FONT_SIZES["body"])
                self.pdf_canvas.drawString(self.margins["left"], self.current_y, f"{category}:")
                
                self.pdf_canvas.setFont(font_regular, self.FONT_SIZES["body"])
                skills_text = ", ".join(skills)
                
                max_width = self.margins["right"] - (self.margins["left"] + 100)
                x = self.margins["left"] + 100
                
                words = skills_text.split()
                lines = []
                current_line = []
                
                for word in words:
                    test_line = ' '.join(current_line + [word])
                    if self.pdf_canvas.stringWidth(test_line, font_regular, self.FONT_SIZES["body"]) <= max_width:
                        current_line.append(word)
                    else:
                        if current_line:
                            lines.append(' '.join(current_line))
                        current_line = [word]
                
                if current_line:
                    lines.append(' '.join(current_line))
                
                if lines:
                    self.pdf_canvas.drawString(x, self.current_y, lines[0])
                    self.current_y -= 15
                    
                    for line in lines[1:]:
                        self.pdf_canvas.drawString(x, self.current_y, line)
                        self.current_y -= 15
                
                self.current_y -= 5
        
        if self.docx_doc:
            for category, skills in skills_dict.items():
                paragraph = self.docx_doc.add_paragraph()
                
                run = paragraph.add_run(f"{category}: ")
                run.font.size = Pt(self.FONT_SIZES["body"])
                run.font.bold = True
                run.font.name = 'Calibri'
                
                run = paragraph.add_run(", ".join(skills))
                run.font.size = Pt(self.FONT_SIZES["body"])
                run.font.name = 'Calibri'
                
                paragraph.paragraph_format.space_after = Pt(6)
    
    def save(self, base_filename: str):
        """Save the document(s)."""
        saved_files = []
        
        if self.pdf_canvas and self.output_format in ["pdf", "both"]:
            pdf_filename = f"{base_filename}.pdf"
            self.pdf_canvas.save()
            saved_files.append(pdf_filename)
            print(f"✓ PDF saved: {pdf_filename}")
        
        if self.docx_doc and self.output_format in ["docx", "both"]:
            docx_filename = f"{base_filename}.docx"
            self.docx_doc.save(docx_filename)
            saved_files.append(docx_filename)
            print(f"✓ DOCX saved: {docx_filename}")
        
        return saved_files
    
    def generate(self, base_filename: str, cv_data: Dict):
        """Generate document(s) from CV data dictionary."""
        # Initialize documents
        if self.output_format in ["pdf", "both"]:
            self._init_pdf(f"{base_filename}.pdf")
        
        if self.output_format in ["docx", "both"]:
            self._init_docx()
        
        # Header
        self.add_name(cv_data.get("name", ""))
        self.add_title(cv_data.get("title", ""))
        
        # Contact info
        contact = cv_data.get("contact", {})
        self.add_contact_info(
            email=contact.get("email", ""),
            phone=contact.get("phone"),
            linkedin=contact.get("linkedin"),
            github=contact.get("github")
        )
        
        # Summary
        if "summary" in cv_data and cv_data["summary"]:
            self.add_section_heading("PROFESSIONAL SUMMARY")
            self.add_paragraph(cv_data["summary"])
        
        # Experience
        if "experience" in cv_data and cv_data["experience"]:
            self.add_section_heading("PROFESSIONAL EXPERIENCE")
            for exp in cv_data["experience"]:
                self.add_job_entry(
                    title=exp.get("title", ""),
                    company=exp.get("company", ""),
                    location=exp.get("location"),
                    dates=exp.get("dates", ""),
                    bullets=exp.get("bullets")
                )
        
        # Education
        if "education" in cv_data and cv_data["education"]:
            self.add_section_heading("EDUCATION")
            for edu in cv_data["education"]:
                self.add_education_entry(
                    degree=edu.get("degree", ""),
                    institution=edu.get("institution", ""),
                    location=edu.get("location"),
                    dates=edu.get("dates", ""),
                    details=edu.get("details")
                )
        
        # Certifications
        if "certifications" in cv_data and cv_data["certifications"]:
            self.add_section_heading("CERTIFICATIONS")
            for cert in cv_data["certifications"]:
                self.add_certification(
                    name=cert.get("name", ""),
                    issuer=cert.get("issuer", ""),
                    date=cert.get("date"),
                    credential_url=cert.get("url")
                )
        
        # Skills
        if "skills" in cv_data and cv_data["skills"]:
            self.add_section_heading("SKILLS")
            self.add_skills_section(cv_data["skills"])
        
        # Save documents
        return self.save(base_filename)
